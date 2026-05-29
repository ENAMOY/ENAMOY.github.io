import docx
import os

print("Start analysis...")
file_path = 'romans/罗马书讲道集 第一部 V2.0 Martha.docx'


try:
    doc = docx.Document(file_path)
    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    
    print("\n--- First 200 paragraphs (truncated) ---")
    for i, p in enumerate(doc.paragraphs[:200]):
        if p.text.strip():
            print(f"Index: {i} | Style: {p.style.name} | Text: {p.text[:50]}")

except Exception as e:
    print(f"Error: {e}")
